"""
Lesson Notes Generation Service
Generates Word documents for each lesson in a course module
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os
from app.services.llm_service import generate_lesson_notes as llm_generate_notes

# Base directory for storing generated notes
NOTES_DIR = os.path.join(os.path.dirname(__file__), 'notes')

def ensure_notes_directory():
    """Create notes directory if it doesn't exist"""
    if not os.path.exists(NOTES_DIR):
        os.makedirs(NOTES_DIR)
    return NOTES_DIR

def sanitize_filename(text):
    """Remove invalid characters from filename"""
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        text = text.replace(char, '_')
    return text.replace(' ', '_').lower()[:50]


def get_existing_notes_file(course_title, module_title, lesson_title):
    """
    Check if notes file already exists for this lesson
    
    Args:
        course_title: Course name
        module_title: Module name
        lesson_title: Lesson title
    
    Returns:
        str: Full file path if exists, None otherwise
    """
    ensure_notes_directory()
    
    # Build expected filename
    filename = f"{sanitize_filename(course_title)}_{sanitize_filename(module_title)}_{sanitize_filename(lesson_title)}.docx"
    file_path = os.path.join(NOTES_DIR, 'notes', filename)
    
    # Check if file exists
    if os.path.exists(file_path):
        return file_path
    
    return None

def generate_lesson_notes(course_title, module_title, lesson, preferences=None):
    """
    Generate comprehensive notes for a lesson using LLM
    
    Args:
        course_title: Title of the course
        module_title: Title of the module containing the lesson
        lesson: Dict containing lesson info (title, summary)
        preferences: Optional learning preferences for personalization
    
    Returns:
        dict: {
            'success': bool,
            'file_path': str or None,
            'message': str,
            'content': str or None
        }
    """
    
    try:
        # Prepare prompt for LLM
        lesson_title = lesson.get('title', 'Untitled Lesson')
        lesson_summary = lesson.get('summary', '')
        
        print(f"Generating notes for: {course_title} > {module_title} > {lesson_title}")
        
        # Call LLM to generate detailed lesson content
        notes_content = llm_generate_notes(
            course_title=course_title,
            module_title=module_title,
            lesson_title=lesson_title,
            lesson_summary=lesson_summary,
            preferences=preferences
        )
        
        if not notes_content:
            return {
                'success': False,
                'file_path': None,
                'message': 'Failed to generate notes content',
                'content': None
            }
        
        # Create Word document
        doc_path = create_notes_document(
            course_title=course_title,
            module_title=module_title,
            lesson_title=lesson_title,
            notes_content=notes_content
        )
        
        if doc_path:
            return {
                'success': True,
                'file_path': doc_path,
                'message': 'Notes generated successfully',
                'content': notes_content
            }
        else:
            return {
                'success': False,
                'file_path': None,
                'message': 'Failed to create document',
                'content': notes_content
            }
            
    except Exception as e:
        print(f"Error generating lesson notes: {str(e)}")
        return {
            'success': False,
            'file_path': None,
            'message': f'Error: {str(e)}',
            'content': None
        }

def create_notes_document(course_title, module_title, lesson_title, notes_content):
    """
    Create a formatted Word document for lesson notes
    
    Args:
        course_title: Course name
        module_title: Module name
        lesson_title: Lesson title
        notes_content: Dict with keys: introduction, key_concepts, examples, summary, additional_resources
    
    Returns:
        str: File path to created document
    """
    
    try:
        # Create new document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Header - Course Title
        heading = doc.add_heading(course_title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        run.font.bold = True
        
        # Module and Lesson info
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Module {module_title}\nLesson: {lesson_title}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.font.italic = True
        
        # Add horizontal line
        doc.add_paragraph('_' * 70)
        
        # Introduction
        if 'introduction' in notes_content:
            doc.add_heading('Introduction', level=2)
            doc.add_paragraph(notes_content['introduction'])
        
        # Key Concepts
        if 'key_concepts' in notes_content and notes_content['key_concepts']:
            doc.add_heading('Key Concepts', level=2)
            for concept in notes_content['key_concepts']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(concept).bold = True
                
                # Add explanation if available
                if isinstance(concept, dict) and 'explanation' in concept:
                    doc.add_paragraph(concept['explanation'])
        
        # Detailed Explanation
        if 'detailed_explanation' in notes_content:
            doc.add_heading('Detailed Explanation', level=2)
            doc.add_paragraph(notes_content['detailed_explanation'])
        
        # Examples
        if 'examples' in notes_content and notes_content['examples']:
            doc.add_heading('Examples', level=2)
            for i, example in enumerate(notes_content['examples'], 1):
                doc.add_paragraph(f'Example {i}:', style='Heading 3')
                doc.add_paragraph(example)
        
        # Practice Exercises
        if 'practice_exercises' in notes_content and notes_content['practice_exercises']:
            doc.add_heading('Practice Exercises', level=2)
            for i, exercise in enumerate(notes_content['practice_exercises'], 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(exercise)
        
        # Summary
        if 'summary' in notes_content:
            doc.add_heading('Summary', level=2)
            summary_para = doc.add_paragraph(notes_content['summary'])
            summary_para.runs[0].font.bold = True
        
        # Additional Resources
        if 'additional_resources' in notes_content and notes_content['additional_resources']:
            doc.add_heading('Additional Resources', level=2)
            for resource in notes_content['additional_resources']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(resource)
        
        # Footer with timestamp
        doc.add_paragraph('_' * 70)
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)
        run.font.italic = True
        
        # Save document
        ensure_notes_directory()
        filename = f"{sanitize_filename(course_title)}_{sanitize_filename(module_title)}_{sanitize_filename(lesson_title)}.docx"
        file_path = os.path.join(NOTES_DIR, filename)
        
        doc.save(file_path)
        print(f"Notes document saved: {file_path}")
        
        return file_path
        
    except Exception as e:
        print(f"Error creating Word document: {str(e)}")
        return None


def get_all_notes_for_course(course_title):
    """
    Get all generated notes files for a course
    
    Args:
        course_title: Title of the course
    
    Returns:
        list: List of dicts with note file info
    """
    notes = []
    course_prefix = sanitize_filename(course_title)
    
    if os.path.exists(NOTES_DIR):
        for filename in os.listdir(NOTES_DIR):
            if filename.startswith(course_prefix) and filename.endswith('.docx'):
                file_path = os.path.join(NOTES_DIR, filename)
                notes.append({
                    'filename': filename,
                    'file_path': file_path,
                    'created_at': datetime.fromtimestamp(os.path.getctime(file_path))
                })
    
    return sorted(notes, key=lambda x: x['created_at'], reverse=True)


def download_notes(file_path):
    """
    Prepare notes file for download
    
    Args:
        file_path: Path to the notes document
    
    Returns:
        tuple: (filename, file_content) or (None, None) if error
    """
    try:
        if os.path.exists(file_path):
            filename = os.path.basename(file_path)
            with open(file_path, 'rb') as f:
                content = f.read()
            return filename, content
        return None, None
    except Exception as e:
        print(f"Error reading notes file: {str(e)}")
        return None, None
